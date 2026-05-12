"""
utils/resume_parser.py — Extracts raw text from PDF and DOCX files.
Uses PyMuPDF (fitz) for PDF and python-docx for DOCX.
"""

import io
from pathlib import Path
from core.logger import get_logger

logger = get_logger(__name__)


def parse_pdf(file_bytes: bytes) -> str:
    """
    Extract text from a PDF file given its raw bytes.

    Args:
        file_bytes: Raw bytes of the PDF file.

    Returns:
        Extracted text string (pages joined with double newline).
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError("PyMuPDF is required for PDF parsing. Run: pip install PyMuPDF")

    text_parts = []
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page_num, page in enumerate(doc):
            page_text = page.get_text("text")
            if page_text.strip():
                text_parts.append(page_text)
            logger.debug("PDF page %d extracted (%d chars)", page_num + 1, len(page_text))

    full_text = "\n\n".join(text_parts)
    logger.info("PDF parsed: %d pages, %d chars total", len(text_parts), len(full_text))
    return full_text


def parse_docx(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX file given its raw bytes.

    Args:
        file_bytes: Raw bytes of the DOCX file.

    Returns:
        Extracted text string (paragraphs joined with newline).
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required. Run: pip install python-docx")

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    paragraphs.append(cell_text)

    full_text = "\n".join(paragraphs)
    logger.info("DOCX parsed: %d paragraphs, %d chars total", len(paragraphs), len(full_text))
    return full_text


def parse_file(filename: str, file_bytes: bytes) -> str:
    """
    Route to the correct parser based on file extension.

    Args:
        filename: Original filename (used for extension detection).
        file_bytes: Raw bytes of the uploaded file.

    Returns:
        Extracted text string.

    Raises:
        ValueError: If the file type is unsupported.
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return parse_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return parse_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: '{ext}'. Supported: .pdf, .docx")
